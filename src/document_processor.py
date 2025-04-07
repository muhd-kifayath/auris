import os
from pathlib import Path
from collections import Counter
import pdfplumber
import cv2
import pytesseract
from docx import Document
import shutil
import re
import fitz

class DocumentProcessor:
    def __init__(self, output_folder="extracted_images"):
        self.output_folder = output_folder
        self.header_footer_threshold = 3
        os.makedirs(self.output_folder, exist_ok=True)

    def clean_text(self, text):
        # Replace weird joins: "endofsentence.THENew" → "end of sentence. THE New"
        text = re.sub(r'(?<=[a-z])(?=[A-Z])', ' ', text)
        # Ensure space after punctuation if missing
        text = re.sub(r'(?<=[.!?])(?=\w)', ' ', text)
        return text


    def extract_text_from_word(self, docx_path):
        """Extracts text from a Word document and replaces images with placeholders."""
        doc = Document(docx_path)
        text = ""

        for para in doc.paragraphs:
            text += para.text + "\n"
            for rel in para.part.rels:
                if "image" in para.part.rels[rel].target_ref:
                    text += "\n[IMAGE PLACEHOLDER]\n"

        return text


    def extract_text_from_pdf(self, pdf_path):
        doc = fitz.open(pdf_path)
        all_pages_text = []
        image_placeholder = "\n[IMAGE PLACEHOLDER]\n"

        for page in doc:
            text_lines = []

            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                for line in block.get("lines", []):
                    line_text = ""
                    for span in line.get("spans", []):
                        span_text = span["text"].strip()
                        font = span.get("font", "")
                        if not span_text:
                            continue
                        # Bold and Italic formatting
                        if "Bold" in font:
                            span_text = f"**{span_text}**"
                        elif "Italic" in font or "Oblique" in font:
                            span_text = f"*{span_text}*"
                        line_text += span_text + " "
                    text_lines.append(line_text.strip())

            page_text = "\n".join(text_lines).strip()
            if page.get_images():
                page_text += image_placeholder
            all_pages_text.append(page_text)

        # Remove repeated headers and footers
        headers = self._find_repeated_lines(all_pages_text, position="top")
        footers = self._find_repeated_lines(all_pages_text, position="bottom")

        cleaned_pages = []
        for page_text in all_pages_text:
            lines = page_text.split("\n")
            if lines and lines[0] in headers:
                lines = lines[1:]
            if lines and lines[-1] in footers:
                lines = lines[:-1]
            cleaned_pages.append("\n".join(lines).strip())

        return "\n\n".join(cleaned_pages)

    def _find_repeated_lines(self, pages_text, position="top"):
        candidates = []
        for page in pages_text:
            lines = page.strip().split("\n")
            if not lines:
                continue
            if position == "top":
                candidates.append(lines[0])
            elif position == "bottom":
                candidates.append(lines[-1])

        common = Counter(candidates).most_common()
        repeated = [line for line, count in common if count >= self.header_footer_threshold]
        return repeated


    def extract_text_from_image(self, image_path):
        """Extracts text from an image using OCR."""
        img = cv2.imread(image_path)
        if img is None:
            raise FileNotFoundError(f"Image file {image_path} not found or could not be read.")
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        return pytesseract.image_to_string(gray)

    def extract_images_from_pdf(self, pdf_path):
        """Extracts images from a PDF and saves them to the output folder."""
        with pdfplumber.open(pdf_path) as pdf:
            for i, page in enumerate(pdf.pages):
                for j, img in enumerate(page.images):
                    image_data = img["stream"].get_data()
                    with open(f"{self.output_folder}/pdf_page_{i}_img_{j}.png", "wb") as f:
                        f.write(image_data)

    def extract_images_from_word(self, docx_path):
        """Extracts images from a Word document and saves them to the output folder."""
        doc = Document(docx_path)
        for i, rel in enumerate(doc.part.rels):
            if "image" in doc.part.rels[rel].target_ref:
                image_data = doc.part.rels[rel].target_part.blob
                with open(f"{self.output_folder}/word_image_{i}.png", "wb") as f:
                    f.write(image_data)


    # def chunk_text(self, text, chunk_size=512):
    #     """Splits text into smaller chunks (approx. chunk_size tokens)."""
    #     paragraphs = re.split(r'\n\s*\n', text)  # Split by paragraphs
    #     chunks = []
    #     current_chunk = []

    #     for paragraph in paragraphs:
    #         words = paragraph.split()
    #         if sum(len(p.split()) for p in current_chunk) + len(words) > chunk_size:
    #             chunks.append("\n".join(current_chunk))
    #             current_chunk = []
    #         current_chunk.append(paragraph)

    #     if current_chunk:
    #         chunks.append("\n".join(current_chunk))

    #     return chunks

    def chunk_text(self, text):
        """Split text based on the smallest markdown-style heading or paragraphs."""

        # Match any heading level (e.g., #, ##, ###...) followed by a space and text
        heading_pattern = re.compile(r'(?=^#{1,6} .*$)', re.MULTILINE)
        headings = heading_pattern.findall(text)

        if headings:
            # Determine the smallest heading level present
            heading_levels = [len(h.split(' ')[0]) for h in headings]
            min_heading_level = max(heading_levels)  # Split using the most specific heading
            print(f"Splitting text using heading level: {min_heading_level}")
            smallest_heading_pattern = re.compile(rf'(?=^#{{{min_heading_level}}} .*$)', re.MULTILINE)
            chunks = [chunk.strip() for chunk in smallest_heading_pattern.split(text) if chunk.strip()]
        else:
            # Fallback to splitting by paragraphs
            chunks = [para.strip() for para in text.split("\n\n") if para.strip()]

        return chunks


    
    def process_document(self, file_path):
        """Processes a document (PDF/Word), extracts text, and replaces images with extracted text.
        Splits text into chunks for better vector storage and retrieval.
        """
        file_extension = Path(file_path).suffix.lower()

        if file_extension == ".pdf":
            text = self.extract_text_from_pdf(file_path)
            # self.extract_images_from_pdf(file_path)
        elif file_extension == ".docx":
            text = self.extract_text_from_word(file_path)
            # self.extract_images_from_word(file_path)
        else:
            raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")

        # Extract text from images and replace placeholders
        extracted_texts = []
        # for image_file in sorted(Path(self.output_folder).glob("*.png")):
        #     extracted_texts.append(self.extract_text_from_image(str(image_file)))

        # for extracted_text in extracted_texts:
        #     text = text.replace("[IMAGE PLACEHOLDER]", extracted_text, 1)

        # Cleanup extracted images
        # shutil.rmtree(self.output_folder)

        # Chunk the text into meaningful sections
        chunks = self.chunk_text(text)

        for i, chunk in enumerate(chunks):
            # Clean the chunked text
            if "[IMAGE PLACEHOLDER]" in chunk:
                chunk = chunk.replace("[IMAGE PLACEHOLDER]", "")

        return text  # Now returns a list of text chunks


# # # 📝 Example Usage
# processor = DocumentProcessor()

# pdf_file_path = "documents/file.pdf"
# # # docx_file_path = "documents/file.docx"

# # Extracted text from PDF
# pdf_text = processor.process_document(pdf_file_path)
# print("Extracted Text from PDF:\n", pdf_text)

# # # Extracted text from Word document
# # # docx_text = processor.process_document(docx_file_path)
# # n = 1
# for i in range(len(pdf_text)):
#     print(f"Chunked Text {i}: {pdf_text[i]}\n")
# # print("Extracted Text from Word:\n", docx_text)
